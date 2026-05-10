"""一次性运行：生成测试 fixture（不进 git 历史外）"""
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

OUT = Path(__file__).parent
OUT.mkdir(parents=True, exist_ok=True)

# 注册中文字体（CID 内置，无需外部字体文件）
pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def make_pdf():
    c = canvas.Canvas(str(OUT / "sample.pdf"), pagesize=A4)
    c.setFont("STSong-Light", 12)
    c.drawString(50, 800, "XX 政务服务平台业务需求")
    c.drawString(50, 770, "3.2.1 门户首页")
    c.drawString(50, 750, "新闻列表 + 政务服务图标")
    c.drawString(50, 720, "3.2.2 新闻管理")
    c.drawString(50, 700, "新增、修改、删除、查看、撤回")
    c.drawString(50, 670, "3.2.3 用户管理")
    c.save()


def make_docx():
    from docx import Document
    d = Document()
    d.add_heading("XX 政务服务平台业务需求", level=1)
    d.add_heading("3.2 业务功能概括", level=2)
    d.add_paragraph("3.2.1 门户首页：新闻列表 + 政务服务图标")
    d.add_paragraph("3.2.2 新闻管理：新增/修改/删除/查看/撤回")
    # 加一个表格（FP 清单常见形式）
    t = d.add_table(rows=3, cols=3)
    t.rows[0].cells[0].text = "模块"
    t.rows[0].cells[1].text = "功能项"
    t.rows[0].cells[2].text = "说明"
    t.rows[1].cells[0].text = "门户"
    t.rows[1].cells[1].text = "首页"
    t.rows[1].cells[2].text = "展示新闻"
    t.rows[2].cells[0].text = "新闻管理"
    t.rows[2].cells[1].text = "新闻新增"
    t.rows[2].cells[2].text = "录入新闻"
    d.save(str(OUT / "sample.docx"))


def make_xlsx():
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "功能清单"
    ws.append(["子系统", "一级模块", "二级模块", "功能项描述"])
    ws.append(["政务服务平台", "业务功能", "门户首页", "新闻列表"])
    ws.append(["政务服务平台", "业务功能", "新闻管理", "新增"])
    ws.append(["政务服务平台", "业务功能", "新闻管理", "修改"])
    ws.append(["政务服务平台", "业务功能", "用户管理", "登录"])
    wb.save(str(OUT / "sample.xlsx"))


if __name__ == "__main__":
    make_pdf()
    make_docx()
    make_xlsx()
    print("✓ fixtures created")
