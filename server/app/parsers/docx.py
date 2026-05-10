from pathlib import Path
import docx as python_docx
from .pdf import ParsedDocument


def parse_docx(path: Path) -> ParsedDocument:
    """提取 Word 文档全文（含段落 + 表格）"""
    if not path.exists():
        raise FileNotFoundError(f"DOCX 文件不存在: {path}")
    d = python_docx.Document(str(path))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            line = " | ".join(c.text.strip() for c in row.cells)
            if line:
                parts.append(line)
    return ParsedDocument(
        text="\n".join(parts),
        page_count=1,  # docx 无页面概念，统一记 1
        metadata={"source": str(path), "type": "docx"})
